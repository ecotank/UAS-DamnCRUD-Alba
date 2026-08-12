import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/><w:left w:val="single" w:sz="18" w:space="0" w:color="1F4E78"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/><w:right w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    
    doc.add_paragraph()

def create_jawaban_docx(output_path):
    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Document Header Table
    header_tbl = doc.add_table(rows=1, cols=1)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_cell = header_tbl.cell(0, 0)
    set_cell_background(h_cell, "1F4E78")
    set_cell_margins(h_cell, top=180, bottom=180, left=200, right=200)

    hp = h_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hrun1 = hp.add_run("LEMBAR JAWABAN UJIAN AKHIR SEMESTER (UAS)\n")
    hrun1.font.name = 'Calibri'
    hrun1.font.size = Pt(16)
    hrun1.font.bold = True
    hrun1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    hrun2 = hp.add_run("POLITEKNIK SIBER DAN SANDI NEGARA\nTAHUN AKADEMIK 2025/2026")
    hrun2.font.name = 'Calibri'
    hrun2.font.size = Pt(12)
    hrun2.font.bold = True
    hrun2.font.color.rgb = RGBColor(0xD9, 0xE1, 0xF2)

    doc.add_paragraph()

    # Metadata Table
    meta_tbl = doc.add_table(rows=5, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("MATA KULIAH", ": Pengujian Perangkat Lunak"),
        ("KELAS / PRODI", ": 4 RPLK (Rekayasa Perangkat Lunak Kripto)"),
        ("HARI, TANGGAL", ": Senin, 12 Agustus 2026"),
        ("DOSEN PENGAMPU", ": Herman Kabetta"),
        ("SIFAT UJIAN", ": Open Book (Laptop)")
    ]

    for idx, (label, val) in enumerate(meta_data):
        row = meta_tbl.rows[idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(2.2)
        c2.width = Inches(4.3)
        set_cell_margins(c1, top=40, bottom=40, left=80, right=80)
        set_cell_margins(c2, top=40, bottom=40, left=80, right=80)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(10)
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(val)
        r2.font.size = Pt(10)

    doc.add_paragraph()

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    # SOAL 1
    add_heading_1("SOAL 1: PERGESERAN PARADIGMA CI/CD DALAM PENGEMBANGAN SISTEM [BOBOT 15%]")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("Continuous Integration (CI) dan Continuous Delivery/Deployment (CD) mengubah pergeseran paradigma (paradigm shift) dari pendekatan tradisional (Waterfall / Manual SDLC) menuju pendekatan modern (Agile / DevOps / DevSecOps).")

    add_heading_2("Tabel Komparasi Paradigma Tradisional vs CI/CD Pipeline:")
    s1_tbl = doc.add_table(rows=6, cols=3)
    s1_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    s1_headers = ["Aspek Pengujian", "Paradigma Lama (Manual/Tradisional)", "Paradigma Baru (CI/CD Pipeline)"]
    s1_widths = [Inches(1.8), Inches(2.3), Inches(2.4)]

    hdr_cells = s1_tbl.rows[0].cells
    for i, title in enumerate(s1_headers):
        hdr_cells[i].width = s1_widths[i]
        set_cell_background(hdr_cells[i], "1F4E78")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)

    s1_rows_data = [
        ("Integrasi Kode", "Big Bang Integration di akhir siklus sprint/rilis. Rentan memicu merge hell dan konflik antar modul.", "Continuous Integration: Push berkali-kali sehari, otomatis diintegrasikan dan diuji via build runner."),
        ("Pengujian (Testing)", "Manual & Late Testing: QA menguji manual di akhir siklus pengembangan.", "Automated & Shift-Left Testing: Pengujian unit, integrasi, dan E2E berjalan otomatis tiap commit."),
        ("Waktu Umpan Balik", "Sangat lambat (butuh hari/minggu untuk menerima laporan bug dari tim QA).", "Sangat cepat (hitungan menit setelah commit dikirim ke pipeline CI)."),
        ("Rilis Perangkat Lunak", "Monolithic Release dalam skala besar beberapa bulan sekali. Risiko kegagalan tinggi.", "Continuous Delivery/Deployment: Perubahan kecil dirilis secara bertahap dan teratur."),
        ("Quality Assurance", "QA menjadi bottleneck karena pengujian regresi manual yang berulang.", "QA berfokus pada eksplorasi skenario kompleks, otomasi, dan pencegahan cacat.")
    ]

    for r_idx, (col1, col2, col3) in enumerate(s1_rows_data, start=1):
        row_cells = s1_tbl.rows[r_idx].cells
        bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        vals = [col1, col2, col3]
        for c_idx in range(3):
            cell = row_cells[c_idx]
            cell.width = s1_widths[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(vals[c_idx])
            run.font.size = Pt(9)

    doc.add_paragraph()

    # SOAL 2
    add_heading_1("SOAL 2: MATRIKS PURE FUNCTIONAL TEST CASE (DamnCRUD) [BOBOT 25%]")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Berikut adalah 50 Pure Functional Test Cases (Functional Testing, Boundary Value & Input Validation Matrix) pada aplikasi DamnCRUD:")

    tc_tbl = doc.add_table(rows=51, cols=6)
    tc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc_headers = ["ID", "TEST CASE OBJECTIVE", "TEST CASE DESCRIPTION", "EXPECTED RESULT", "ACTUAL RESULT", "STATUS"]
    tc_widths = [Inches(0.6), Inches(1.3), Inches(1.6), Inches(1.5), Inches(1.3), Inches(0.7)]

    thdr_cells = tc_tbl.rows[0].cells
    for i, title in enumerate(tc_headers):
        thdr_cells[i].width = tc_widths[i]
        set_cell_background(thdr_cells[i], "1F4E78")
        set_cell_margins(thdr_cells[i], top=100, bottom=100, left=60, right=60)
        p = thdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8.5)

    test_cases_data = [
        ("TC-001", "Autentikasi User Valid", "Input username 'admin' dan password 'nimda666!'", "Sistem mengautentikasi dan redirect ke index.php", "Masuk Dashboard, tampil 'Howdy admin!'", "PASS"),
        ("TC-002", "Autentikasi Password Invalid", "Input username 'admin' dan password 'wrongpass'", "Menolak autentikasi dan tampilkan pesan error", "Tampil 'Damn, wrong credentials!!'", "PASS"),
        ("TC-003", "Autentikasi Username Invalid", "Input username 'user_palsu' dan password 'nimda666!'", "Menolak autentikasi dan tampilkan pesan error", "Tampil 'Damn, wrong credentials!!'", "PASS"),
        ("TC-004", "Validasi Required Form Login", "Menekan submit tanpa mengisi form login", "Browser memicu validasi atribut required", "Submit ditahan dengan notifikasi browser", "PASS"),
        ("TC-005", "Fungsi Logout Sesi", "Menekan tombol Sign out pada header", "Sesi dihancurkan (session_destroy()) & redirect login", "Logout sukses dan redirect ke login.php", "PASS"),
        ("TC-006", "Proteksi Sesi Dashboard", "Akses langsung index.php tanpa sesi login", "Menolak akses dan redirect ke login.php", "Otomatis di-redirect ke login.php", "PASS"),
        ("TC-007", "Proteksi Sesi Profil", "Akses langsung profil.php tanpa sesi login", "Menolak akses dan redirect ke login.php", "Otomatis di-redirect ke login.php", "PASS"),
        ("TC-008", "Proteksi Sesi VPage", "Akses langsung vpage.php tanpa sesi login", "Menolak akses dan redirect ke login.php", "Otomatis di-redirect ke login.php", "PASS"),
        ("TC-009", "Proteksi Sesi Create Form", "Akses langsung create.php tanpa sesi login", "Menolak akses dan redirect ke login.php", "Otomatis di-redirect ke login.php", "PASS"),
        ("TC-010", "Proteksi Sesi Update Form", "Akses langsung update.php?id=1 tanpa sesi login", "Menolak akses dan redirect ke login.php", "Otomatis di-redirect ke login.php", "PASS"),
        ("TC-011", "Proteksi Sesi Delete Action", "Akses langsung delete.php?id=1 tanpa sesi login", "Menolak akses dan redirect ke login.php", "Otomatis di-redirect ke login.php", "PASS"),
        ("TC-012", "Tambah Kontak Baru Valid", "Isi form create.php data valid & tekan Save", "Data tersimpan ke DB & muncul di Dashboard", "Kontak tersimpan & tampil di tabel", "PASS"),
        ("TC-013", "Tambah Kontak Nama Kosong", "Mengosongkan field Name di create.php", "Browser memblokir submit atribut required", "Form menolak submit validasi browser", "PASS"),
        ("TC-014", "Tambah Kontak Email Kosong", "Mengosongkan field Email di create.php", "Browser memblokir submit atribut required", "Form menolak submit validasi browser", "PASS"),
        ("TC-015", "Tambah Kontak Phone Kosong", "Mengosongkan field Phone di create.php", "Browser memblokir submit atribut required", "Form menolak submit validasi browser", "PASS"),
        ("TC-016", "Tambah Kontak Title Kosong", "Mengosongkan field Title di create.php", "Browser memblokir submit atribut required", "Form menolak submit validasi browser", "PASS"),
        ("TC-017", "Navigation Cancel Create", "Tekan tombol Cancel pada form create.php", "Form dibatalkan & redirect ke index.php", "Redirect kembali ke Dashboard", "PASS"),
        ("TC-018", "Penanganan Variabel $id", "Kirim form create.php kode asli", "Berhasil insert data tanpa error PHP", "Terdeteksi bug $id undefined (Fixed)", "PASS"),
        ("TC-019", "Syntax Email Tanpa @", "Input email 'budi.example.com' tanpa @", "Browser/form memblokir format email salah", "Form ter-submit (HTML type='text')", "FAIL"),
        ("TC-020", "Validasi Phone Karakter Abjad", "Input string 'ABCD-EFGH' pada field Phone", "Aplikasi memvalidasi hanya format angka", "Form menerima string abjad (type='text')", "FAIL"),
        ("TC-021", "Input Whitespace Nama", "Input string spasi '   ' pada Name", "Sistem trim() atau tolak nama kosong", "Kontak tersimpan nama spasi kosong", "FAIL"),
        ("TC-022", "Boundary String Input Nama", "Input string Nama sepanjang 300 karakter", "DB memotong string max 255 karakter aman", "MariaDB memotong string 255 char", "PASS"),
        ("TC-023", "Special Characters & Emoji", "Input nama 'Budi @#$% 😃' di create.php", "DB utf8mb4 menyimpan karakter tanpa corrupt", "Simbol & emoji ter-render presisi", "PASS"),
        ("TC-024", "Tampilan Tabel Kontak", "Navigasi ke Dashboard index.php", "Menampilkan tabel kontak dari database", "Seluruh kolom kontak tampil rapi", "PASS"),
        ("TC-025", "Fitur Search DataTables", "Ketik nama 'John' pada input search", "Tabel memfilter baris yang cocok secara instan", "DataTables menyaring baris real-time", "PASS"),
        ("TC-026", "Navigasi Pagination", "Tekan tombol halaman 2 / Next DataTables", "Tabel berpindah menampilkan baris 11-15", "Halaman berpindah menampilkan data baru", "PASS"),
        ("TC-027", "Sorting Kolom Tabel", "Klik header kolom Name / Created", "Baris terurut Ascending (A-Z) / Descending", "Urutan baris berubah sesuai sorting", "PASS"),
        ("TC-028", "Filter Wildcards Table", "Input wildcard '%' atau '*' di DataTables", "DataTables handle string pencarian literal", "DataTables menyaring string literal", "PASS"),
        ("TC-029", "Buka Form Update ID Valid", "Klik edit pada kontak ID 1", "Form update.php?id=1 terbuka terisi data", "Form terbuka menampilkan data ID 1", "PASS"),
        ("TC-030", "Update ID Tidak Ada", "Akses URL update.php?id=99999", "Hentikan eksekusi 'Contact doesn't exist!'", "Tampil 'Contact doesn't exist!'", "PASS"),
        ("TC-031", "Update Tanpa Parameter ID", "Akses URL update.php tanpa ?id=", "Hentikan eksekusi 'No ID specified!'", "Tampil 'No ID specified!'", "PASS"),
        ("TC-032", "Simpan Update Kontak", "Ubah Name & Title pada update.php?id=1", "Data ter-update di DB & redirect Dashboard", "Data ter-update & redirect index.php", "PASS"),
        ("TC-033", "Field Phone Form Update", "Buka form update.php?id=1 versi asli", "Field Phone memuat nomor telepon lama", "Terdeteksi bug value='' kosong (Fixed)", "PASS"),
        ("TC-034", "Navigation Cancel Update", "Tekan tombol Cancel pada update.php", "Form dibatalkan & redirect ke index.php", "Redirect ke Dashboard tanpa ubah data", "PASS"),
        ("TC-035", "Parameter ID Non-Numeric", "Akses update.php?id=abc string abjad", "PDO cast ID aman & tampilkan error data", "Tampil 'Contact doesn't exist!'", "PASS"),
        ("TC-036", "Parameter ID Negatif", "Akses update.php?id=-5 ID negatif", "SQL handle ID < 0 secara aman", "Tampil 'Contact doesn't exist!'", "PASS"),
        ("TC-037", "Parameter ID Desimal Float", "Akses update.php?id=1.5 desimal float", "Query PDO convert ID integer 1 aman", "Data ID 1 tampil tanpa SQL error", "PASS"),
        ("TC-038", "Hapus Kontak Alert OK", "Klik delete & tekan OK pada confirm alert", "Record terhapus dari DB & baris hilang", "Alert terkonfirmasi & data terhapus", "PASS"),
        ("TC-039", "Hapus Kontak Alert Cancel", "Klik delete & tekan Cancel pada alert", "Tindakan hapus dibatalkan, data tetap ada", "Alert dibatalkan, data tidak terhapus", "PASS"),
        ("TC-040", "Delete Non-Existent ID", "Akses delete.php?id=99999 ID tidak ada", "DELETE 0 baris & redirect index.php", "Redirect index.php tanpa error PHP", "PASS"),
        ("TC-041", "Direct Delete Tanpa ID", "Akses delete.php tanpa parameter ?id=", "Hentikan eksekusi 'No ID specified!'", "Tampil 'No ID specified!'", "PASS"),
        ("TC-042", "Tampilan Halaman Profil", "Navigasi ke profil.php", "Menampilkan profil, form upload, & user info", "Halaman profil terbuka sempurna", "PASS"),
        ("TC-043", "Upload Foto Profil JPG", "Pilih file .jpg & tekan tombol Change", "File terunggah ke image/profile.jpg", "Foto terunggah tanpa error", "PASS"),
        ("TC-044", "Upload Foto Profil JPEG", "Pilih file .jpeg & tekan tombol Change", "File terunggah ke image/profile.jpg", "Foto terunggah tanpa error", "PASS"),
        ("TC-045", "Upload Profil Format PNG", "Pilih file .png & tekan tombol Change", "Tolak file & tampilkan pesan error ekstensi", "Tampil 'Ekstensi tidak diijinkan...'", "PASS"),
        ("TC-046", "Upload Profil Format PDF", "Pilih file .pdf & tekan tombol Change", "Tolak file & tampilkan pesan error ekstensi", "Tampil 'Ekstensi tidak diijinkan...'", "PASS"),
        ("TC-047", "Upload Ekstensi Uppercase", "Upload file 'AVATAR.JPG' huruf besar", "Sistem strtolower() & izinkan file .JPG", "Ext check case-sensitive, file ditolak", "FAIL"),
        ("TC-048", "Upload Double Extension", "Upload file 'avatar.jpg.png' di profil", "Sistem periksa ekstensi akhir (.png) & tolak", "Ekstensi .png terdeteksi & ditolak", "PASS"),
        ("TC-049", "Upload File Kosong 0-Byte", "Upload file 0-byte ber-ekstensi .jpg", "Sistem periksa filesize > 0 & tolak", "File 0-byte terunggah tanpa validasi", "FAIL"),
        ("TC-050", "Submit Upload Tanpa File", "Tekan Change tanpa memilih file", "Form reloaded tanpa melakukan upload", "Halaman reloaded tanpa perubahan", "PASS")
    ]

    for r_idx, row_data in enumerate(test_cases_data, start=1):
        r_cells = tc_tbl.rows[r_idx].cells
        bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx in range(6):
            cell = r_cells[c_idx]
            cell.width = tc_widths[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=60, bottom=60, left=50, right=50)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.1
            
            val = row_data[c_idx]
            run = p.add_run(val)
            run.font.size = Pt(8)
            
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.bold = True
            elif c_idx == 5:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.bold = True
                if val == "PASS":
                    run.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)
                else:
                    run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)

    doc.add_paragraph()

    # SOAL 3
    add_heading_1("SOAL 3: LANGKAH-LANGKAH & SKRIP OTOMASI SELENIUM PYTHON [BOBOT 30%]")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Pengujian fungsional diotomatisasi menggunakan Python 3.11, Selenium WebDriver (Headless Chrome), dan Pytest. Lima test case terpilih diuji secara terintegrasi:")

    add_heading_2("Skrip Fixture Pytest (tests/conftest.py):")
    conftest_code = """import pytest
import os
import time
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager

BASE_URL = os.environ.get("BASE_URL", "http://127.0.0.1:8000")
SELENIUM_HUB = os.environ.get("SELENIUM_HUB", None)

@pytest.fixture(scope="function")
def driver():
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--remote-allow-origins=*")
    chrome_options.add_argument("--window-size=1920,1080")

    if SELENIUM_HUB:
        driver = webdriver.Remote(command_executor=SELENIUM_HUB, options=chrome_options)
    else:
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=chrome_options)

    driver.implicitly_wait(5)
    yield driver
    driver.quit()

@pytest.fixture(scope="function")
def logged_in_driver(driver):
    target_url = BASE_URL.rstrip('/')
    driver.get(f"{target_url}/login.php")
    driver.find_element("id", "inputUsername").send_keys("admin")
    driver.find_element("id", "inputPassword").send_keys("nimda666!")
    driver.find_element("xpath", "//button[@type='submit']").click()
    time.sleep(1)
    assert "Dashboard" in driver.title or "Howdy" in driver.page_source
    return driver"""
    add_code_block(doc, conftest_code)

    add_heading_2("Skrip Selenium Functional Test (tests/test_damncrud.py):")
    test_code = """import pytest
import os
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.common.alert import Alert
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

BASE_URL = os.environ.get("BASE_URL", "http://host.docker.internal:8000")

def test_tc003_add_new_contact(logged_in_driver):
    driver = logged_in_driver
    driver.find_element(By.CLASS_NAME, "create-contact").click()
    time.sleep(1)
    assert "Add new contact" in driver.title
    
    driver.find_element(By.ID, "name").send_keys("Budi Santoso Automation")
    driver.find_element(By.ID, "email").send_keys("budi.auto@example.com")
    driver.find_element(By.ID, "phone").send_keys("081299887766")
    driver.find_element(By.ID, "title").send_keys("Cryptographer Tester")
    driver.find_element(By.XPATH, "//input[@type='submit']").click()
    time.sleep(1)
    
    assert "Dashboard" in driver.title
    search_input = driver.find_element(By.XPATH, "//input[@type='search']")
    search_input.send_keys("Budi Santoso Automation")
    time.sleep(0.5)
    assert "Budi Santoso Automation" in driver.page_source

def test_tc004_update_contact(logged_in_driver):
    driver = logged_in_driver
    driver.find_element(By.XPATH, "//a[contains(@href, 'update.php')]").click()
    time.sleep(1)
    
    name_field = driver.find_element(By.ID, "name")
    name_field.clear()
    name_field.send_keys("John Does Updated")
    driver.find_element(By.XPATH, "//input[@type='submit']").click()
    time.sleep(1)
    
    assert "Dashboard" in driver.title
    assert "John Does Updated" in driver.page_source

def test_tc005_delete_contact(logged_in_driver):
    driver = logged_in_driver
    info_before = driver.find_element(By.ID, "employee_info").text
    driver.find_element(By.XPATH, "//a[contains(@href, 'delete.php')]").click()
    time.sleep(0.5)
    
    try:
        WebDriverWait(driver, 3).until(EC.alert_is_present())
        alert = driver.switch_to.alert
        alert.accept()
        time.sleep(1)
    except Exception:
        pass
    
    driver.get(f"{BASE_URL.rstrip('/')}/index.php")
    time.sleep(1)
    info_after = driver.find_element(By.ID, "employee_info").text
    assert info_before != info_after

def test_tc006_upload_profile_image(logged_in_driver, tmp_path):
    driver = logged_in_driver
    driver.find_element(By.XPATH, "//a[contains(@href, 'profil.php')]").click()
    time.sleep(1)
    
    sample_jpg = tmp_path / "test_avatar.jpg"
    sample_jpg.write_bytes(b"\\xFF\\xD8\\xFF\\xE0\\x00\\x10JFIF\\x00\\x01\\x01\\x01\\x00`\\x00`\\x00\\x00\\xFF\\xD9")
    
    driver.find_element(By.ID, "formFile").send_keys(str(sample_jpg))
    driver.find_element(By.XPATH, "//button[@type='submit']").click()
    time.sleep(1)
    assert "Profil" in driver.title
    assert "Ekstensi tidak diijinkan" not in driver.page_source

def test_tc007_vpage_functional_submission(logged_in_driver):
    driver = logged_in_driver
    driver.find_element(By.XPATH, "//a[contains(@href, 'vpage.php')]").click()
    time.sleep(1)
    assert "Dummy Page XSS Detect" in driver.page_source
    
    test_text = "Testing Functional Output"
    driver.find_element(By.NAME, "thing").send_keys(test_text)
    driver.find_element(By.NAME, "submit").click()
    time.sleep(1)
    assert f"Your thing is {test_text}" in driver.page_source"""
    add_code_block(doc, test_code)

    # SOAL 4
    add_heading_1("SOAL 4: PIPELINE CI/CD GITHUB ACTIONS & PARALLEL TESTING PYTEST [BOBOT 30%]")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Pipeline CI/CD dikonfigurasi menggunakan GitHub Actions (.github/workflows/ci.yml). Pengujian otomatis dijalankan secara paralel menggunakan modul pytest-xdist (pytest -n auto) di dalam service container terisolasi:")

    add_heading_2("Workflow GitHub Actions (.github/workflows/ci.yml):")
    ci_code = """name: DamnCRUD CI/CD Pipeline

on:
  push:
    branches: [ main, master, uas2026 ]
  pull_request:
    branches: [ main, master, uas2026 ]

jobs:
  functional-tests:
    name: Parallel Functional Testing (Pytest + Selenium)
    runs-on: ubuntu-latest

    services:
      mysql:
        image: mysql:8.0
        env:
          MYSQL_ROOT_PASSWORD: root123
          MYSQL_DATABASE: badcrud
        ports:
          - 3306:3306
        options: >-
          --health-cmd="mysqladmin ping"
          --health-interval=10s
          --health-timeout=5s
          --health-retries=3

    steps:
      - name: Checkout Repository Code
        uses: actions/checkout@v4

      - name: Setup PHP Environment
        uses: shivammathur/setup-php@v2
        with:
          php-version: '8.2'
          extensions: pdo, pdo_mysql, mysqli

      - name: Initialize & Import Database
        run: |
          mysql -h 127.0.0.1 -u root -proot123 -e "CREATE DATABASE IF NOT EXISTS damncrud;"
          mysql -h 127.0.0.1 -u root -proot123 badcrud < db/damncrud.sql
          mysql -h 127.0.0.1 -u root -proot123 damncrud < db/damncrud.sql

      - name: Start PHP Built-in Server
        run: |
          php -S 127.0.0.1:8000 &
          sleep 2

      - name: Setup Python Environment
        uses: actions/setup-python@v5
        with:
          python-version: '3.11'

      - name: Install Testing Dependencies
        run: |
          python -m pip install --upgrade pip
          pip install pytest pytest-xdist selenium webdriver-manager

      - name: Run Pytest Parallel Functional Tests
        env:
          BASE_URL: http://127.0.0.1:8000
        run: |
          pytest tests/test_damncrud.py -v -n auto"""
    add_code_block(doc, ci_code)

    add_heading_2("Bukti Link Repositori GitHub Submission:")
    p_link = doc.add_paragraph()
    run_link = p_link.add_run("https://github.com/ecotank/UAS-DamnCRUD-Alba")
    run_link.font.bold = True
    run_link.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    run_link.font.underline = True

    doc.save(output_path)
    print(f"File DOCX berhasil dibuat di: {output_path}")

if __name__ == "__main__":
    out = "d:/alba/Akademik/Semester 6/Pengujian Perangkat Lunak/UAS/Jawaban_UAS_Pengujian_Perangkat_Lunak_v2.docx"
    create_jawaban_docx(out)
